import streamlit as st
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime, date
import re
import os
import csv
import pandas as pd
from zoneinfo import ZoneInfo
import socket
import requests
PATIENT_DB_FILE = "patient_database.csv"

def parse_date(date_string):
    if pd.isna(date_string) or date_string == 'N/A' or date_string == '':
        return pd.NaT
    try:
        return pd.to_datetime(date_string)
    except:
        return pd.NaT


def load_patient_database():
    columns = [
        "Rut", "Nombre", "Edad", "Sexo", "Domicilio", "Fecha de ingreso", "Diagnóstico",
        "Alergias", "Tabaquismo", "Medicamentos", "Antiagregantes plaquetarios", "Anticoagulantes",
        "Antecedentes mórbidos", "Temperatura", "Presión arterial", "Saturación O2",
        "Anamnesis", "Examen físico", "Escala de Glasgow", "Hemiparesia", "Paraparesia",
        "Focalidad", "Exámenes","Plan", "Reposo", "Tromboprofilaxis farmacológica", "Hidratación",
        "Régimen nutricional", "Equipo multidisciplinario", "Antibiótico 1",
        "Fecha de inicio Antibiotico 1", "Días de antibiótico 1", "Antibiótico 2",
        "Fecha de inicio Antibiotico 2", "Días de antibiótico 2", "Retiro sonda foley",
        "Retiro de CVC", "Curación por enfermería", "Instalación sonda nasogástrica",
        "Oxigenoterapia", "Hemoglucotest", "Precauciones", "Exámenes de laboratorio",
        "Firma médico"
    ]

    if not os.path.exists(PATIENT_DB_FILE):
        return pd.DataFrame(columns=columns)
    try:
        df = pd.read_csv(PATIENT_DB_FILE)
        date_columns = ["Fecha de ingreso", "Fecha de inicio Antibiotico 1", "Fecha de inicio Antibiotico 2"]
        for col in date_columns:
            df[col] = df[col].apply(parse_date)
        return df
    except Exception as e:
        st.error(f"Error loading patient database: {str(e)}")
        return pd.DataFrame(columns=columns)


def save_patient_database(df):
    try:
        # Convert date columns to string before saving
        date_columns = ["Fecha de ingreso", "Fecha de inicio Antibiotico 1", "Fecha de inicio Antibiotico 2"]
        for col in date_columns:
            df[col] = df[col].astype(str).replace('NaT', 'N/A')

        df.to_csv(PATIENT_DB_FILE, index=False)
    except Exception as e:
        st.error(f"Error saving patient database: {str(e)}")
def lookup_patient(rut, df):
    patient = df[df["Rut"] == rut]
    if not patient.empty:
        return patient.iloc[0].to_dict()
    return None


def add_patient(data, df):
    # Convert date fields to datetime objects
    date_fields = ["Fecha de ingreso", "Fecha de inicio Antibiotico 1", "Fecha de inicio Antibiotico 2"]
    for field in date_fields:
        if field in data:
            data[field] = parse_date(data[field])

    new_patient = pd.DataFrame([data])
    df = df[df["Rut"] != data["Rut"]]  # Remove existing entry if any
    df = pd.concat([df, new_patient], ignore_index=True)
    save_patient_database(df)
    return df


def save_dict_to_csv(data_dict, filename=None):
    """
    Save a dictionary to a CSV file.

    :param data_dict: Dictionary containing the data to be saved
    :param filename: Optional; If not provided, a filename will be generated based on the current datetime
    :return: The filename of the saved CSV file, or None if saving failed
    """
    try:
        if not filename:
            filename = f"patient_data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"

        # Get the absolute path
        abs_filename = os.path.abspath(filename)
        print(f"Attempting to save CSV file: {abs_filename}")

        # Ensure the directory exists
        os.makedirs(os.path.dirname(abs_filename), exist_ok=True)

        # Get all unique keys from the data dictionary
        fieldnames = list(data_dict.keys())

        # Write data to CSV file
        with open(abs_filename, mode='w', newline='', encoding='utf-8') as file:
            writer = csv.DictWriter(file, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerow(data_dict)

        print(f"CSV file saved successfully: {abs_filename}")
        return abs_filename
    except Exception as e:
        print(f"Error saving CSV file: {e}")
        return None

def create_word_document(data):
    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = section.right_margin = Inches(0.5)
    section.top_margin = section.bottom_margin = Inches(0.3)
    section.orientation = WD_ORIENT.PORTRAIT

    # Set default paragraph format
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(9)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.space_after = Pt(0)

    ## Add title
    chile_tz = ZoneInfo("America/Santiago")  # Chile is in GMT-3
    now = datetime.now(chile_tz)
    current_time = now.strftime("%H:%M")
    title = doc.add_paragraph("Evolución médica neurocirugía")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(9)
    title.runs[0].font.bold = True

    subtitle = doc.add_paragraph(f"{data['Nombre']} , {data['Rut']} , fecha: {data['Fecha']},hora: {current_time}\n")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(9)

    # Define sections and their corresponding fields
    sections = {
        "Información del Paciente": ["Sexo","Edad","Fecha de ingreso", "Días de hospitalización","Alergias","Domicilio"],
        "Diagnóstico": ["Diagnostico", "Antecedentes mórbidos"],
        "Evaluación clínica": ["Temperatura", "Presión arterial", "Saturación O2","Examen físico","Escala de Glasgow", "Hemiparesia","Paraparesia","Focalidad","Exámenes"],
        "Tratamiento": ["Reposo", "Tromboprofilaxis farmacológica", "Régimen nutricional", "Hidratación","Equipo multidisciplinario"],
        "Indicaciones enfermería": ["Retiro sonda foley", "Retiro de CVC", "Curación por enfermería", "Instalación sonda nasogástrica", "Oxigenoterapia", "Exámenes de laboratorio", "Hemoglucotest","Precauciones"],
        "Firma médico": ["Firma médico"]
    }

    for section_title, fields in sections.items():
        # Add section title
        doc.add_paragraph().add_run(section_title).bold = True

        # Create a 2-column table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths
        table.columns[0].width = Inches(3.75)
        table.columns[1].width = Inches(3.75)

        row_cells = table.rows[0].cells

        for i, field in enumerate(fields):
            if i % 2 == 0 and i > 0:
                # Add a new row for every two fields
                row_cells = table.add_row().cells

            cell = row_cells[i % 2]
            p = cell.paragraphs[0]
            p.add_run(f"{field}: ").bold = True
            if field in data:
                # Format the admission date without time
                if field == "Fecha de ingreso":
                    admission_date = data[field]
                    if isinstance(admission_date, str):
                        try:
                            admission_date = datetime.strptime(admission_date, "%d-%m-%Y")
                        except ValueError:
                            # If parsing fails, it might already be a datetime object
                            pass
                    p.add_run(admission_date.strftime("%d-%m-%Y"))
                else:
                    p.add_run(str(data[field]))
        # Apply smaller font size to table contents
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

        # Add space after table
        doc.add_paragraph()
        # For Información del Paciente, add an additional row for medications
        if section_title == "Información del Paciente":
            row_cells = table.add_row().cells

            # First column: Medicamentos
            cell_med = row_cells[0]
            p_med = cell_med.paragraphs[0]
            run = p_med.add_run("Medicamentos crónicos: ")
            run.bold = True
            run.font.size = Pt(8)
            p_med.add_run(data.get('Medicamentos', 'No registrados')).font.size = Pt(8)

            # Second column: Antiagregantes and Anticoagulantes
            cell_anti = row_cells[1]
            p_anti = cell_anti.paragraphs[0]

            run = p_anti.add_run("Antiagregantes plaquetarios: ")
            run.bold = True
            run.font.size = Pt(8)
            p_anti.add_run(f"{data.get('Antiagregantes plaquetarios', 'No')}").font.size = Pt(8)

            p_anti.add_run("\n")  # Add a newline between antiagregantes and anticoagulantes

            run = p_anti.add_run("Anticoagulantes: ")
            run.bold = True
            run.font.size = Pt(8)
            p_anti.add_run(f"{data.get('Anticoagulantes', 'No')}").font.size = Pt(8)

            # Remove extra space after the paragraphs
            p_med.paragraph_format.space_after = Pt(0)
            p_anti.paragraph_format.space_after = Pt(0)

            # Apply smaller font size to table contents
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
                    # Remove extra space after each paragraph in the table
                    paragraph.paragraph_format.space_after = Pt(0)
        if section_title == "Evolución clínica":
            doc.add_paragraph().add_run("Anamnesis").bold = True
            anamnesis_table = doc.add_table(rows=1, cols=1)
            anamnesis_table.style = 'Table Grid'
            anamnesis_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            anamnesis_table.columns[0].width = Inches(7.5)  # Full width of the page

            anamnesis_cell = anamnesis_table.cell(0, 0)
            anamnesis_paragraph = anamnesis_cell.paragraphs[0]
            anamnesis_paragraph.add_run(data.get("Anamnesis", ""))

            # Apply smaller font size to anamnesis table contents
            for paragraph in anamnesis_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

            # Add space after anamnesis table
            doc.add_paragraph()

        # Add separate table for Plan after the Tratamiento section
        if section_title == "Tratamiento":
            doc.add_paragraph().add_run("Plan").bold = True
            plan_table = doc.add_table(rows=1, cols=1)
            plan_table.style = 'Table Grid'
            plan_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            plan_table.columns[0].width = Inches(7.5)  # Full width of the page

            plan_cell = plan_table.cell(0, 0)
            plan_paragraph = plan_cell.paragraphs[0]
            plan_paragraph.add_run(data.get("Plan", ""))

            # Apply smaller font size to plan table contents
            for paragraph in plan_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

            # Add space after plan table
            doc.add_paragraph()
            # Add antibiotic information only if it's present
            if data.get("Antibiótico 1") != "Ninguno" or data.get("Antibiótico 2") != "Ninguno":
                doc.add_paragraph().add_run("Antibióticos").bold = True
                atb_table = doc.add_table(rows=1, cols=2)
                atb_table.style = 'Table Grid'
                atb_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                atb_table.columns[0].width = Inches(3.75)
                atb_table.columns[1].width = Inches(3.75)

                row_cells = atb_table.rows[0].cells

                # Antibiótico 1
                if data.get("Antibiótico 1") != "Ninguno":
                    cell = row_cells[0]
                    p = cell.paragraphs[0]
                    p.add_run("Antibiótico 1: ").bold = True
                    p.add_run(f"{data.get('Antibiótico 1', '')}")
                    p.add_run("\nFecha de inicio: ").bold = True
                    p.add_run(f"{data.get('Fecha de inicio Antibiotico 1', '')}")
                    p.add_run("\nDías de tratamiento: ").bold = True
                    p.add_run(f"{data.get('Días de antibiótico 1', '')}")

                # Antibiótico 2
                if data.get("Antibiótico 2") != "Ninguno":
                    cell = row_cells[1] if data.get("Antibiótico 1") != "Ninguno" else row_cells[0]
                    p = cell.paragraphs[0]
                    p.add_run("Antibiótico 2: ").bold = True
                    p.add_run(f"{data.get('Antibiótico 2', '')}")
                    p.add_run("\nFecha de inicio: ").bold = True
                    p.add_run(f"{data.get('Fecha de inicio Antibiotico 2', '')}")
                    p.add_run("\nDías de tratamiento: ").bold = True
                    p.add_run(f"{data.get('Días de antibiótico 2', '')}")

                # Apply smaller font size to antibiotic table contents
                for row in atb_table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(6)

                # Add space after antibiotic table
                doc.add_paragraph()
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_text = f"Unidad de neurocirugía, Hospital de Curicó"
    footer_paragraph.text = footer_text
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_paragraph.runs[0]
    footer_run.font.size = Pt(6)
    footer_run.font.name = 'Arial'
    # Create a filename-safe version of the patient's name
    safe_name = re.sub(r'[^\w\-_\. ]', '_', data['Nombre'])
    safe_name = safe_name.replace(' ', '_')
    filename = f"{safe_name}_{datetime.now().strftime('%d%m%Y_%H%M')}.docx"
    doc.save(filename)
    return filename


def validate_form(data):
    required_fields = [
        "Nombre", "Rut", "Edad", "Sexo", "Domicilio", "Fecha de ingreso",
        "Alergias", "Tabaquismo", "Medicamentos", "Antiagregantes plaquetarios",
        "Anticoagulantes", "Temperatura", "Presión arterial", "Saturación O2",
        "Anamnesis", "Examen físico", "Escala de Glasgow", "Focalidad", "Diagnostico",
        "Plan", "Reposo", "Tromboprofilaxis farmacológica", "Hidratación",
        "Régimen nutricional", "Firma médico"
    ]

    missing_fields = [field for field in required_fields if not data.get(field)]

    if missing_fields:
        st.error(f"Por favor, complete los siguientes campos obligatorios: {', '.join(missing_fields)}")
        return False
    return True


def main():
    st.set_page_config(page_title="Evolución médica", layout="wide")
    st.title("Evolución médica neurocirugía")

    # Use Chile time zone for current date
    chile_tz = ZoneInfo("America/Santiago")
    current_date = st.date_input("Fecha actual", value=datetime.now(chile_tz).date())

    # Load patient database
    patient_df = load_patient_database()

    # Initialize patient_info
    patient_info = {}

    # Patient Information section
    st.subheader("Información del Paciente")
    rut = st.text_input("Rut")

    if rut:
        patient_info = lookup_patient(rut, patient_df) or {}
        if patient_info:
            name = st.text_input("Nombre", value=patient_info.get("Nombre", ""), disabled=True)
            age = st.number_input("Edad", value=patient_info.get("Edad", 0), disabled=True)
            gender = st.selectbox("Sexo", ["Masculino", "Femenino"],
                                  index=["Masculino", "Femenino"].index(patient_info.get("Sexo", "Masculino")))
            domicilio = st.selectbox("Domicilio",
                                     ["Curicó", "Molina", "Sagrada Familia", 'Romeral', 'Hualañe', 'Licantén',
                                      'Rauco', 'Teno', 'Vichuquén', 'Otro'],
                                     index=["Curicó", "Molina", "Sagrada Familia", 'Romeral', 'Hualañe', 'Licantén',
                                            'Rauco', 'Teno', 'Vichuquén', 'Otro'].index(
                                         patient_info.get("Domicilio", "Curicó")))
            admission_date = st.date_input("Fecha de ingreso",
                                           value=pd.to_datetime(
                                               patient_info.get("Fecha de ingreso")).date() if pd.notna(
                                               patient_info.get("Fecha de ingreso")) else None)
        else:
            st.warning("Paciente no encontrado. Por favor, ingrese la información manualmente.")
            name = st.text_input("Nombre")
            age = st.number_input("Edad", min_value=0, max_value=120)
            gender = st.selectbox("Sexo", ["Masculino", "Femenino"])
            domicilio = st.selectbox("Domicilio",
                                     ["Curicó", "Molina", "Sagrada Familia", 'Romeral', 'Hualañe', 'Licantén',
                                      'Rauco', 'Teno', 'Vichuquén', 'Otro'])
            admission_date = st.date_input("Fecha de ingreso")
    else:
        name = st.text_input("Nombre")
        age = st.number_input("Edad", min_value=0, max_value=120)
        gender = st.selectbox("Sexo", ["Masculino", "Femenino"])
        domicilio = st.selectbox("Domicilio",
                                 ["Curicó", "Molina", "Sagrada Familia", 'Romeral', 'Hualañe', 'Licantén', 'Rauco',
                                  'Teno', 'Vichuquén', 'Otro'])
        admission_date = st.date_input("Fecha de ingreso")

    # Medical Details section
    st.subheader("Antecedentes médicos")
    col1, col2 = st.columns(2)
    with col1:
        alergias = st.text_input("Alergias", value=patient_info.get("Alergias", ""))
        tabaquismo = st.selectbox("Tabaquismo", ["No", "Si"],
                                  index=["No", "Si"].index(patient_info.get("Tabaquismo", "No")))
        fármacos = st.text_input("Medicamentos crónicos", value=patient_info.get("Medicamentos", ""))
        aspirina = st.selectbox("Antiagregantes plaquetarios", ["No", "Si"],
                                index=["No", "Si"].index(patient_info.get("Antiagregantes plaquetarios", "No")))
        taco = st.selectbox("Anticoagulantes", ["No", "Si"],
                            index=["No", "Si"].index(patient_info.get("Anticoagulantes", "No")))

    with col2:
        st.write("Antecedentes mórbidos")
        morbidos_options = ["Diabetes Mellitus NIR", "Diabetes Mellitus IR", "HTA", "Hipotiroidismo",
                            "Enfermedad renal crónica", "EPOC", "Asma Bronquial", "Daño hepático crónico",
                            "Cardiopatía Coronaria", "Insuficiencia cardíaca", "Arritmia"]
        morbidos_selections = {}
        saved_morbidos = patient_info.get("Antecedentes mórbidos", "").split(", ") if patient_info.get(
            "Antecedentes mórbidos") else []
        for option in morbidos_options:
            morbidos_selections[option] = st.checkbox(option, value=option in saved_morbidos)

    st.subheader("Evaluación clínica")
    col1, col2, col3 = st.columns(3)
    with col1:
        temp = st.text_input("Temperatura (grados)")
    with col2:
        blood_pressure = st.text_input("Presión arterial (sistólica/diastólica)")
    with col3:
        sat02 = st.text_input("Saturación O2")
    medical_history = st.text_area("Anamnesis")
    examen_fisico = st.text_area("Exámen físico")
    st.write("Exámen neurológico")
    st.write("Escala de Glasgow")
    col1, col2, col3 = st.columns(3)
    with col1:
        ocular_options = ["O1", "O2", "O3", "O4"]
        ocular_selections = {}
        for option in ocular_options:
            ocular_selections[option] = st.checkbox(option)
    with col2:
        verbal_options = ["V1", "V2", "V3", "V4", "V5"]
        verbal_selections = {}
        for option in verbal_options:
            verbal_selections[option] = st.checkbox(option)
    with col3:
        motor_options = ["M1", "M2", "M3", 'M4', 'M5', "M6"]
        motor_selections = {}
        for option in motor_options:
            motor_selections[option] = st.checkbox(option)

    col1, col2, col3 = st.columns(3)
    with col1:
        hemiparesia_txt=st.selectbox("Hemiparesia", ["No", "Derecha", "Izquierda"])
        hemiparesia_options = ["MRC  0 sin contracción visible", "MRC 1 contracción visible", "MRC 2 no vence gravedad", "MRC 3 vence gravedad", "MRC 4 vence resistencia","MRC 5 normal"]
        hemiparesia_selections = {}
        for option in hemiparesia_options:
            hemiparesia_selections[option] = st.checkbox(option)
    with col2:
        paraparesia_txt=st.selectbox("Paraparesia", ["No", "Si","Tetraparesia"])
        paraparesia_options = ["MRC0  sin contracción visible","MRC1 contracción visible", "MRC2 no vence gravedad", "MRC3 vence gravedad", "MRC4 vence resistencia", "MRC5 normal"]
        paraparesia_selections = {}
        for option in paraparesia_options:
            paraparesia_selections[option] = st.checkbox(option)
    with col3:
        focalidad=st.text_input("Focalidad neurológica:")
    examenes = st.text_area("Exámenes")

    # Additional Details section
    st.subheader("Diagnóstico")
    diagnostico = st.text_area("Diagnóstico")
    st.subheader("Tratamiento")
    plan = st.text_area("Plan")
    col1, col2, col3 = st.columns(3)
    with col1:
        reposo = st.selectbox("Reposo",
                              ["Absoluto cero grados", 'Absoluto cabecera en 30 grados', 'Absoluto semisentado',
                               "Levantar asisitdo", "Relativo"])
        trombo = st.selectbox("Tromboprofilaxis farmacológica", ["No", "Si"])
        suero= st.text_input("Hidratación  (ml/lr)")
    with col2:
        # Régimen section (nested within col2)
        st.write("Régimen nutricional")
        regimen_options = ["Ayunas", "Hídrico", "Líquidos","Licuado blando", "Liviano", "Común", "Hiposódico", "Diabético", "Enteral"]
        regimen_selections = {}
        for option in regimen_options:
            regimen_selections[option] = st.checkbox(option)
    with col3:
        st.write("Equipo multidisciplinario")
        equipo_options = ["Kinesioterapia motora ", "Kinesioterapia respiratoria", "Fonoaudiología",
                          "Terapia Ocupacional", "Fisiatría"]
        equipo_selections = {}
        for option in equipo_options:
            equipo_selections[option] = st.checkbox(option)
    atb1 = st.selectbox("Antibiótico 1", ["Ninguno", "Cefazolina", "Cloxacilina", "Ceftriaxona", "Vancomicina", "Metronidazol", "Tazonam", "Gentamicina", "Ciprofloxacino", "Levofloxacino", "Cotrimoxazol", "Meropenem"])
    date_atb1 = st.date_input("Fecha inicio ATB 1", value=None)
    atb2 = st.selectbox("Antibiótico 2", ["Ninguno", "Cefazolina", "Cloxacilina", "Ceftriaxona", "Vancomicina", "Metronidazol", "Tazonam", "Gentamicina", "Ciprofloxacino", "Levofloxacino", "Cotrimoxazol", "Meropenem"])
    date_atb2 = st.date_input("Fecha inicio ATB 2", value=None)
    st.subheader("Indicaciones enfermería")
    col1, col2, col3 = st.columns(3)
    with col1:
        foley = st.selectbox("Retiro sonda foley", ['No', "Si"])
        cvc = st.selectbox("Retiro de CVC", ["No", "Si"])
    with col2:
        curacion = st.selectbox("Curación por enfermería", ["No", "Si"])
        SNG = st.selectbox("Instalación sonda nasogástrica", ["No", "Si"])
        precauciones=st.selectbox("Precauciones", ["No", "Gotas","Aéreo","Contacto"])

    with col3:
        oxigeno = st.selectbox("Oxigenoterapia", ["No", "Bigotera", 'Mascarilla'])
        examenes = st.selectbox("Exámenes de laboratorio",
                                ["No", "Urgencia (orden amarilla)", 'Laboratorio (orden blanca)'])
        HGT = st.selectbox("Hemoglucotest", ["No", "Cada 6 hrs", "Cada 6 hrs + insulina cristalina según esquema"])
    st.subheader("Firma médico")
    firma = st.selectbox("Neurocirujano",
                         ["Dr. Nicolás González Romo", "Dr. Patricio Giménez Hermosilla", "Dr.José Villamediana","Dr.Héctor Aceituno"])


    if st.button("Guardar"):
        # Process selections
        selected_regimens = [option for option, selected in regimen_selections.items() if selected]
        regimen_str = ", ".join(selected_regimens) if selected_regimens else "Ninguno seleccionado"
        selected_equipo = [option for option, selected in equipo_selections.items() if selected]
        equipo_str = ", ".join(selected_equipo) if selected_equipo else "Ninguno seleccionado"
        selected_morbidos = [option for option, selected in morbidos_selections.items() if selected]
        morbidos_str = ", ".join(selected_morbidos) if selected_morbidos else "Ninguno seleccionado"
        selected_hemiparesia = [option for option, selected in hemiparesia_selections.items() if selected]
        hemiparesia_str = f"{hemiparesia_txt}/" + ", ".join(selected_hemiparesia) if selected_hemiparesia else "Ninguno seleccionado"
        selected_paraparesia = [option for option, selected in paraparesia_selections.items() if selected]
        paraparesia_str = f"{paraparesia_txt}/" +", ".join(selected_paraparesia) if selected_paraparesia else "Ninguno seleccionado"
        # Process Glasgow scale
        ocular_score = next((key for key, value in ocular_selections.items() if value), "N/A")
        verbal_score = next((key for key, value in verbal_selections.items() if value), "N/A")
        motor_score = next((key for key, value in motor_selections.items() if value), "N/A")
        glasgow_score = f"Ocular: {ocular_score}, Verbal: {verbal_score}, Motor: {motor_score}"

        # Calculate hospitalization time
        hospitalization_time = (current_date - admission_date).days

        # Calculate antibiotic therapy duration
        atb1_time = (current_date - date_atb1).days + 1 if atb1 != "Ninguno" and date_atb1 is not None else 0
        atb2_time = (current_date - date_atb2).days + 1 if atb2 != "Ninguno" and date_atb2 is not None else 0

        data = {
            "Nombre": name,
            "Rut": rut,
            "Edad": age,
            "Sexo": gender,
            "Domicilio": domicilio,
            "Fecha": current_date.strftime("%d-%m-%Y"),
            "Fecha de ingreso": admission_date.strftime("%d-%m-%Y"),
            "Días de hospitalización": f"{hospitalization_time} días",
            "Alergias": alergias,
            "Tabaquismo": tabaquismo,
            "Medicamentos": fármacos,
            "Antiagregantes plaquetarios": aspirina,
            "Anticoagulantes": taco,
            "Antecedentes mórbidos": morbidos_str,
            "Temperatura": f"{temp} grados",
            "Presión arterial": blood_pressure,
            "Saturación O2": sat02,
            "Anamnesis": medical_history,
            "Examen físico": examen_fisico,
            "Escala de Glasgow": glasgow_score,
            "Hemiparesia": hemiparesia_str,
            "Paraparesia": paraparesia_str,
            "Focalidad": focalidad,
            "Exámenes": examenes,
            "Diagnostico": diagnostico,
            "Plan": plan,
            "Reposo": reposo,
            "Tromboprofilaxis farmacológica": trombo,
            "Hidratación": suero,
            "Régimen nutricional": regimen_str,
            "Equipo multidisciplinario": equipo_str,
            "Antibiótico 1": atb1,
            "Fecha de inicio Antibiotico 1": date_atb1.strftime("%d-%m-%Y") if date_atb1 else "N/A",
            "Días de antibiótico 1": f"{atb1_time} días" if atb1_time > 0 else "N/A",
            "Antibiótico 2": atb2,
            "Fecha de inicio Antibiotico 2": date_atb2.strftime("%d-%m-%Y") if date_atb2 else "N/A",
            "Días de antibiótico 2":f"{atb2_time} días" if atb2_time > 0 else "N/A",
            "Retiro sonda foley": foley,
            "Retiro de CVC": cvc,
            "Curación por enfermería": curacion,
            "Instalación sonda nasogástrica": SNG,
            "Oxigenoterapia": oxigeno,
            "Hemoglucotest": HGT,
            "Precauciones": precauciones,
            "Exámenes de laboratorio": examenes,
            "Firma médico": firma
        }
        def ensure_reports_folder():
            if not os.path.exists("reports"):
                os.makedirs("reports")

        if validate_form(data):
            patient_df = add_patient(data, patient_df)

            ensure_reports_folder()
            filename = create_word_document(data)
            new_filename = os.path.join("reports", os.path.basename(filename))
            os.rename(filename, new_filename)

            csv_filename = save_dict_to_csv(data)
            st.success(f"Archivo guardado: {new_filename}")
            st.download_button(
                label="Descargar documento Word",
                data=open(new_filename, "rb"),
                file_name=os.path.basename(new_filename),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.warning("Por favor, complete todos los campos obligatorios antes de guardar.")

if __name__ == "__main__":
    main()