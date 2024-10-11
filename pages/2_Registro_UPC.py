import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from datetime import datetime
import pandas as pd

PATIENT_DB_FILE = "patient_database.csv"


def load_patient_database():
    columns = ["Nombre", "RUT", "Edad", "Diagnóstico", "Evaluación Clínica",
               "Ventilación Mecánica", "Drogas Vasoactivas", "Nivel de Sedación (SAS)",
               "Evaluación Pupilar", "Examen Motor", "Herida Quirúrgica",
               "Escala de Coma de Glasgow", "Exámenes de Laboratorio",
               "Estudios de Imagen", "Plan de Tratamiento"]
    if not os.path.exists(PATIENT_DB_FILE):
        return pd.DataFrame(columns=columns)
    return pd.read_csv(PATIENT_DB_FILE)


def save_patient_database(df):
    df.to_csv(PATIENT_DB_FILE, index=False)


def create_word_document(data):
    doc = Document()

    # Set default paragraph format
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.space_after = Pt(0)

    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = section.right_margin = Inches(1)
    section.top_margin = section.bottom_margin = Inches(1)

    title = doc.add_paragraph("Evolución médica neurocirugía UPC")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(14)
    title.runs[0].font.bold = True

    current_time = datetime.now().strftime("%d-%m-%Y %H:%M")
    subtitle = doc.add_paragraph(f"Fecha y hora: {current_time}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(10)

    # Patient Information Table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(5.0)

    row = table.rows[0].cells
    row[0].text = "Nombre"
    row[1].text = data['Nombre']
    row = table.rows[1].cells
    row[0].text = "RUT"
    row[1].text = data['RUT']
    row = table.rows[2].cells
    row[0].text = "Edad"
    row[1].text = str(data['Edad'])
    row = table.rows[3].cells
    row[0].text = "Diagnóstico"
    row[1].text = data['Diagnóstico']

    doc.add_paragraph()
    doc.add_paragraph().add_run("Evaluación Clínica").bold = True

    # General Assessment Table
    general_assessment_table = doc.add_table(rows=1, cols=1)
    general_assessment_table.style = 'Table Grid'
    general_assessment_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    general_assessment_table.columns[0].width = Inches(6.5)
    general_assessment_table.rows[0].cells[0].text = data['Evaluación Clínica']

    doc.add_paragraph()
    doc.add_paragraph().add_run("Detalles de Evaluación Clínica").bold = True

    # Detailed Clinical Assessment Table
    assessment_table = doc.add_table(rows=6, cols=2)
    assessment_table.style = 'Table Grid'
    assessment_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    assessment_table.columns[0].width = Inches(2.0)
    assessment_table.columns[1].width = Inches(4.5)

    row = assessment_table.rows[0].cells
    row[0].text = "Ventilación Mecánica"
    row[1].text = data['Ventilación Mecánica']
    row = assessment_table.rows[1].cells
    row[0].text = "Drogas Vasoactivas"
    row[1].text = data['Drogas Vasoactivas']
    row = assessment_table.rows[2].cells
    row[0].text = "Nivel de Sedación (SAS)"
    row[1].text = data['Nivel de Sedación (SAS)']
    row = assessment_table.rows[3].cells
    row[0].text = "Evaluación Pupilar"
    row[1].text = data['Evaluación Pupilar']
    row = assessment_table.rows[4].cells
    row[0].text = "Examen Motor"
    row[1].text = data['Examen Motor']
    row = assessment_table.rows[5].cells
    row[0].text = "Herida Quirúrgica"
    row[1].text = data['Herida Quirúrgica']

    doc.add_paragraph()
    doc.add_paragraph().add_run("Escala de Coma de Glasgow").bold = True
    glasgow_table = doc.add_table(rows=1, cols=1)
    glasgow_table.style = 'Table Grid'
    glasgow_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    glasgow_cell = glasgow_table.rows[0].cells[0]
    glasgow_cell.text = data['Escala de Coma de Glasgow']
    glasgow_table.columns[0].width = Inches(6.5)

    doc.add_paragraph()
    doc.add_paragraph().add_run("Estudios Clínicos").bold = True
    studies_table = doc.add_table(rows=2, cols=2)
    studies_table.style = 'Table Grid'
    studies_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    studies_table.columns[0].width = Inches(2.0)
    studies_table.columns[1].width = Inches(4.5)

    row = studies_table.rows[0].cells
    row[0].text = "Exámenes de Laboratorio"
    row[1].text = data['Exámenes de Laboratorio']
    row = studies_table.rows[1].cells
    row[0].text = "Estudios de Imagen"
    row[1].text = data['Estudios de Imagen']

    doc.add_paragraph()
    doc.add_paragraph().add_run("Plan de Tratamiento").bold = True
    plan_table = doc.add_table(rows=1, cols=1)
    plan_table.style = 'Table Grid'
    plan_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    plan_cell = plan_table.rows[0].cells[0]
    plan_cell.text = data['Plan de Tratamiento']
    plan_table.columns[0].width = Inches(6.5)

    # Apply consistent formatting to all paragraphs and table cells
    for paragraph in doc.paragraphs:
        paragraph.style = doc.styles['Normal']
        for run in paragraph.runs:
            run.font.size = Pt(10)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = doc.styles['Normal']
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

    filename = f"Evolucion_medica_neurocirugia_UPC_{data['Nombre']}_{datetime.now().strftime('%d%m%Y_%H%M')}.docx"
    doc.save(filename)
    return filename


def main():
    st.set_page_config(page_title="Evolución médica neurocirugía UPC", layout="wide")
    st.title("Evolución médica neurocirugía UPC")

    patient_df = load_patient_database()

    name = st.text_input("Nombre")
    rut = st.text_input("RUT")
    age = st.number_input("Edad", min_value=0, max_value=120)
    diagnosis = st.text_input("Diagnóstico")
    clinical_assessment = st.text_area("Evaluación Clínica General")

    st.subheader("Detalles de Evaluación Clínica")
    mechanical_ventilation = st.selectbox("Ventilación Mecánica", ["No", "Si"])
    vasoactive_drugs = st.selectbox("Drogas Vasoactivas", ["No", "Si"])
    sedation_level = st.selectbox("Nivel de Sedación (SAS)",
                                  ["1 - No responde",
                                   "2 - Responde solo a estímulos fuertes",
                                   "3 - Difícil de despertar",
                                   "4 - Tranquilo y cooperador",
                                   "5 - Agitado",
                                   "6 - Muy agitado",
                                   "7 - Peligrosamente agitado"])

    pupillary_assessment = st.text_area("Evaluación Pupilar")
    motor_exam = st.text_area("Examen Motor")
    surgical_wound = st.text_area("Evaluación de Herida Quirúrgica")

    glasgow_scale = st.text_input("Escala de Coma de Glasgow")

    st.subheader("Estudios Clínicos")
    lab_tests = st.text_area("Exámenes de Laboratorio")
    imaging_studies = st.text_area("Estudios de Imagen")

    treatment_plan = st.text_area("Plan de Tratamiento")

    if st.button("Guardar"):
        new_patient = pd.DataFrame({
            "Nombre": [name],
            "RUT": [rut],
            "Edad": [age],
            "Diagnóstico": [diagnosis],
            "Evaluación Clínica": [clinical_assessment],
            "Ventilación Mecánica": [mechanical_ventilation],
            "Drogas Vasoactivas": [vasoactive_drugs],
            "Nivel de Sedación (SAS)": [sedation_level],
            "Evaluación Pupilar": [pupillary_assessment],
            "Examen Motor": [motor_exam],
            "Herida Quirúrgica": [surgical_wound],
            "Escala de Coma de Glasgow": [glasgow_scale],
            "Exámenes de Laboratorio": [lab_tests],
            "Estudios de Imagen": [imaging_studies],
            "Plan de Tratamiento": [treatment_plan]
        })

        patient_df = pd.concat([patient_df, new_patient], ignore_index=True)
        save_patient_database(patient_df)

        filename = create_word_document(new_patient.iloc[0].to_dict())

        st.success(f"Registro guardado. Documento creado: {filename}")

        with open(filename, "rb") as file:
            st.download_button(
                label="Descargar documento Word",
                data=file,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )


if __name__ == "__main__":
    main()