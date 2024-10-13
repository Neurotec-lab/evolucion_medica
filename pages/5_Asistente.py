import os
import tempfile
import streamlit as st
from langchain_community.llms import Ollama
from langchain_community.document_loaders import DirectoryLoader
from langchain_community.document_loaders import UnstructuredWordDocumentLoader, PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import OllamaEmbeddings
import numpy as np
from docx import Document
from io import BytesIO

# Initialize Ollama
ollama = Ollama(base_url='http://localhost:11434', model="llama3.1:8b")

# Initialize OllamaEmbeddings
oembed = OllamaEmbeddings(base_url="http://localhost:11434", model="nomic-embed-text")


# Function to load and process documents
@st.cache_resource
def load_and_process_documents(uploaded_files):
    if uploaded_files:
        # Create a temporary directory to store uploaded files
        with tempfile.TemporaryDirectory() as temp_dir:
            # Save uploaded files to the temporary directory
            for uploaded_file in uploaded_files:
                with open(os.path.join(temp_dir, uploaded_file.name), "wb") as f:
                    f.write(uploaded_file.getvalue())

            # Load documents from the temporary directory
            loaders = {
                "docx": UnstructuredWordDocumentLoader,
                "pdf": PyPDFLoader
            }
            loader = DirectoryLoader(temp_dir, glob="*.*",
                                     loader_cls=lambda file_path: loaders[file_path.split(".")[-1]](file_path))
            documents = loader.load()

            # Split the documents into chunks
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=0)
            all_splits = text_splitter.split_documents(documents)

            # Create embeddings
            texts = [doc.page_content for doc in all_splits]
            embeddings = oembed.embed_documents(texts)

            return texts, embeddings
    else:
        return None, None


# Function to perform similarity search
def similarity_search(query, texts, embeddings, k=3):
    query_embedding = oembed.embed_query(query)
    similarities = np.dot(embeddings, query_embedding)
    top_k_indices = np.argsort(similarities)[-k:][::-1]
    return [texts[i] for i in top_k_indices]


# Function to generate clinical summary
def generate_clinical_summary(texts, embeddings, patient_name):
    prompt = f"Please create a clear and comprehensive clinical summary in spanish for patient {patient_name}, including personal data (name, RUT, age, admission date), clinical diagnosis, lab results such as hematocrit, hemoglobin, sodium, white blood counts and creatinin levels,and imaging studies (ct scan and mri reports), as well as surgical and medical treatment. Provide a clear and comprehensive overview of the case, starting from the initial presentation. The use of personal data was authorized by the patient"
    relevant_docs = similarity_search(prompt, texts, embeddings)
    context = "\n".join(relevant_docs)
    full_prompt = f"Context: {context}\n\nTask: {prompt}\n\nSummary:"
    summary = ollama.invoke(full_prompt)
    return summary


# Function to create and download DOCX file
def create_docx(summary, patient_name):
    doc = Document()
    doc.add_heading(f'Clinical Summary - {patient_name}', 0)
    doc.add_paragraph(summary)

    # Save the document to a BytesIO object
    docx_file = BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    return docx_file


# Streamlit UI
st.title('Asistente médico')

# Patient name input
patient_name = st.text_input('Nombre del paciente')

# File upload widget
uploaded_files = st.file_uploader("Subir documentos médicos", type=["docx", "pdf"], accept_multiple_files=True)

# Load and process documents when files are uploaded
texts, embeddings = load_and_process_documents(uploaded_files)

# Create a text input box for the user
prompt = st.text_input('Ingresar solicitud')

if prompt and texts and embeddings:
    # Perform similarity search
    relevant_docs = similarity_search(prompt, texts, embeddings)

    # Prepare context for Ollama
    context = "\n".join(relevant_docs)

    # Generate response using Ollama
    full_prompt = f"Context: {context}\n\nQuestion: {prompt}\n\nAnswer:"
    response = ollama.invoke(full_prompt)

    st.write(response)

# Add "Resumen clínico" button
if st.button("Resumen clínico"):
    if texts and embeddings and patient_name:
        summary = generate_clinical_summary(texts, embeddings, patient_name)
        st.write(summary)

        # Create download button for DOCX
        docx_file = create_docx(summary, patient_name)
        st.download_button(
            label="Descargar resumen clínico (DOCX)",
            data=docx_file,
            file_name=f"resumen_clinico_{patient_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    elif not patient_name:
        st.write("Por favor, ingrese el nombre del paciente antes de generar el resumen clínico.")
    else:
        st.write("Por favor, suba documentos médicos antes de generar el resumen clínico.")

st.markdown("""
    <style>
    .footer {
        position: fixed;
        left: 0;
        bottom: 0;
        width: 100%;
        background-color: #f0f2f6;
        color: #31333F;
        text-align: center;
        padding: 10px 0;
        font-size: 14px;
        border-top: 1px solid #d1d1d1;
    }
    </style>
    <div class="footer">
        <p><strong>AVISO DE CONFIDENCIALIDAD:</strong> La información contenida en este sistema es confidencial y está protegida por ley. El acceso no autorizado está prohibido.</p>
        <p>© 2024 Neurocorp Spa. Todos los derechos reservados.</p>
    </div>
    """, unsafe_allow_html=True)