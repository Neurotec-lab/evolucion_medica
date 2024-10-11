import streamlit as st
import pandas as pd
import os
from io import BytesIO
from datetime import datetime, date
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.enum.table import WD_TABLE_ALIGNMENT


# File path for the patient database
PATIENT_DB_FILE = "../patient_database.csv"


def load_patient_database():
    if not os.path.exists(PATIENT_DB_FILE):
        return pd.DataFrame(
            columns=["Rut", "Nombre", "Edad", "Fecha de ingreso", "Diagnostico", "Plan", "Ubicación", "Estado",
                     "Fecha de alta"])
    df = pd.read_csv(PATIENT_DB_FILE)
    if "Estado" not in df.columns:
        df["Estado"] = "Activo"
    if "Fecha de alta" not in df.columns:
        df["Fecha de alta"] = None
    # Ensure 'Rut' column exists (case-insensitive)
    rut_column = next((col for col in df.columns if col.lower() == 'rut'), None)
    if rut_column is None:
        st.error("Error: La columna 'Rut' no se encuentra en el archivo CSV.")
        return pd.DataFrame()
    if rut_column != 'Rut':
        df = df.rename(columns={rut_column: 'Rut'})
    return df


def save_patient_database(df):
    df.to_csv(PATIENT_DB_FILE, index=False)


def discharge_patient(rut):
    df = load_patient_database()
    df.loc[df["Rut"] == rut, "Estado"] = "Alta"
    df.loc[df["Rut"] == rut, "Fecha de alta"] = datetime.now().strftime("%d-%m-%Y")
    save_patient_database(df)


def update_location(rut, new_location):
    df = load_patient_database()
    df.loc[df["Rut"] == rut, "Ubicación"] = new_location
    save_patient_database(df)


def export_to_csv(df):
    return df.to_csv(index=False).encode('utf-8')


def calculate_hospitalization_days(admission_date):
    if pd.isna(admission_date) or admission_date == '':
        return "N/A"

    try:
        # Parse the date in YYYY-MM-DD format
        parsed_date = datetime.strptime(admission_date, "%Y-%m-%d").date()
        today = date.today()
        days = (today - parsed_date).days
        return days
    except ValueError as e:
        st.error(f"Error parsing date '{admission_date}': {e}")
        return "Formato de fecha inválido"


def export_to_docx(df):
    doc = Document()

    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    doc.add_heading('Listado pacientes hospitalizados Neurocirugía', 0)

    # Add current date and time
    current_datetime = datetime.now().strftime("%d-%m-%Y %H:%M:%S")
    doc.add_paragraph(f"Fecha: {current_datetime}")

    # Add table
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set table width to 100% of page width
    table.width = Inches(7)  # Assuming standard page width of 8.5 inches minus 0.5 inch margins on each side

    hdr_cells = table.rows[0].cells
    headers = ['Rut', 'Nombre', 'Edad', 'Fecha de ingreso', 'Días de Hospitalización', 'Diagnostico', 'Plan',
               'Ubicación']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, column in enumerate(headers):
            if column == 'Días de Hospitalización':
                value = calculate_hospitalization_days(row.get('Fecha de ingreso', "N/A"))
            else:
                value = row.get(column, "N/A")
            if pd.isna(value):
                value = "N/A"
            elif isinstance(value, float):
                if value.is_integer():
                    value = int(value)
                value = str(value)
            row_cells[i].text = str(value)
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save the document to a BytesIO object
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

def import_from_csv(uploaded_file):
    if uploaded_file is not None:
        try:
            new_df = pd.read_csv(uploaded_file)
            df = load_patient_database()

            # Ensure 'Rut' column exists in the uploaded file
            rut_column = next((col for col in new_df.columns if col.lower() == 'rut'), None)
            if rut_column is None:
                st.error("Error: La columna 'Rut' no se encuentra en el archivo CSV cargado.")
                return False
            if rut_column != 'Rut':
                new_df = new_df.rename(columns={rut_column: 'Rut'})

            # Update existing patients and add new ones, avoiding duplicates
            for _, row in new_df.iterrows():
                if row["Rut"] in df["Rut"].values:
                    # Update existing patient
                    df.loc[df["Rut"] == row["Rut"], df.columns.intersection(new_df.columns)] = row
                else:
                    # Add new patient
                    df = pd.concat([df, pd.DataFrame([row])], ignore_index=True)

            save_patient_database(df)
            return True
        except Exception as e:
            st.error(f"Error al cargar el archivo: {e}")
    return False


def reset_to_original_database():
    if os.path.exists(PATIENT_DB_FILE):
        df = pd.read_csv(PATIENT_DB_FILE)
        # Ensure all patients are set to "Activo" and clear "Fecha de alta"
        df["Estado"] = "Activo"
        df["Fecha de alta"] = None
        save_patient_database(df)
        return True
    return False


def main():
    st.set_page_config(page_title="Lista de Pacientes Hospitalizados", layout="wide")
    st.title("Lista de Pacientes Hospitalizados")

    # Use session state to track if reset was clicked
    if 'reset_clicked' not in st.session_state:
        st.session_state.reset_clicked = False

    col1, col2, col3, col4 = st.columns([2, 2, 1, 1])

    with col1:
        # Load CSV button
        uploaded_file = st.file_uploader("Cargar lista de pacientes desde CSV", type="csv")
        if uploaded_file is not None:
            if import_from_csv(uploaded_file):
                st.success("Archivo CSV cargado exitosamente")
                st.rerun()

    with col2:
        # Reset button
        if st.button("Reiniciar desde base de datos original"):
            if reset_to_original_database():
                st.session_state.reset_clicked = True
                st.success("Base de datos reiniciada exitosamente")
                st.rerun()
            else:
                st.error("No se pudo reiniciar la base de datos")

    # Load patient database
    df = load_patient_database()
    if df.empty:
        st.error("No se pudo cargar la base de datos de pacientes.")
        return

    active_df = df[df["Estado"] == "Activo"]

    with col3:
        # Export CSV button
        if not active_df.empty:
            st.download_button(
                label="Exportar a CSV",
                data=export_to_csv(active_df),
                file_name="pacientes_hospitalizados.csv",
                mime="text/csv",
            )

    with col4:
        # Export DOCX button
        if not active_df.empty:
            docx_buffer = export_to_docx(active_df)
            st.download_button(
                label="Exportar a DOCX",
                data=docx_buffer,
                file_name="pacientes_hospitalizados.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    # Display patient table
    if not active_df.empty:
        col1, col2, col3, col4, col5, col6, col7, col8, col9, col10 = st.columns([2, 2, 1, 2, 1, 2, 3, 3, 2, 1])
        col1.write("**Rut**")
        col2.write("**Nombre**")
        col3.write("**Edad**")
        col4.write("**Fecha de ingreso**")
        col5.write("**Días**")
        col6.write("**Ubicación**")
        col7.write("**Diagnóstico**")
        col8.write("**Plan**")
        col9.write("**Actualizar**")
        col10.write("**Alta**")
        for index, row in active_df.iterrows():
            col1, col2, col3, col4, col5, col6, col7, col8, col9, col10 = st.columns([2, 2, 1, 2, 1, 2, 3, 3, 2, 1])
            with col1:
                st.write(str(row["Rut"]))
            with col2:
                st.write(row["Nombre"])
            with col3:
                st.write(str(row["Edad"]))
            with col4:
                st.write(row["Fecha de ingreso"])
            with col5:
                hospitalization_days = calculate_hospitalization_days(row["Fecha de ingreso"])
                st.write(f"{hospitalization_days}" if isinstance(hospitalization_days, int) else hospitalization_days)
            with col6:
                location = st.text_input("Ubicación", value=str(row.get("Ubicación", "")), key=f"location_{row['Rut']}")
                if location != str(row.get("Ubicación", "")):
                    update_location(row["Rut"], location)
            with col7:
                st.write(row["Diagnostico"] if pd.notna(row["Diagnostico"]) else "No especificado")
            with col8:
                st.write(row["Plan"] if pd.notna(row["Plan"]) else "No especificado")
            with col9:
                if st.button("Actualizar", key=f"update_{row['Rut']}"):
                    st.success("Ubicación actualizada")
                    st.rerun()
            with col10:
                if st.button("Alta", key=f"discharge_{row['Rut']}"):
                    discharge_patient(row["Rut"])
                    st.success("Paciente dado de alta")
                    st.rerun()
    else:
        st.write("No hay pacientes hospitalizados en este momento.")

    # Clear the reset flag after displaying the data
    if st.session_state.reset_clicked:
        st.session_state.reset_clicked = False


if __name__ == "__main__":
    main()